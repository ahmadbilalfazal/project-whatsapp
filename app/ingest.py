import os
import io
import PyPDF2
import docx
from app import db
from app.models import Document, Chunk

CHUNK_SIZE = 500
CHUNK_OVERLAP = 100


def _persist_chunks(texts, metadatas):
    for metadata, text in zip(metadatas, texts):
        chunk = Chunk(
            business_id=metadata.get('business_id'),
            doc_id=metadata.get('doc_id'),
            page=metadata.get('page'),
            text=text,
            vector_id=None,
        )
        db.session.add(chunk)
    db.session.commit()


def _get_vector_store():
    try:
        from app.vector_store import VectorStore
        return VectorStore()
    except Exception:
        return None


def _read_docx(file_stream):
    doc = docx.Document(file_stream)
    pages = []
    text = '\n'.join(p.text for p in doc.paragraphs)
    pages.append(text)
    return pages


def _read_pdf(file_stream):
    # Use PyPDF2 to extract text from PDF pages
    reader = PyPDF2.PdfReader(file_stream)
    pages = []
    for p in reader.pages:
        try:
            text = p.extract_text() or ''
        except Exception:
            text = ''
        pages.append(text)
    return pages


def _read_txt(file_stream):
    content = file_stream.read().decode('utf-8', errors='ignore')
    return [content]


def _chunk_text(text):
    chunks = []
    start = 0
    L = len(text)
    while start < L:
        end = min(start + CHUNK_SIZE, L)
        chunk = text[start:end]
        chunks.append(chunk)
        if end >= L:
            break
        start = max(0, end - CHUNK_OVERLAP)
    return chunks


def ingest_file(stream, filename, business_id):
    ext = os.path.splitext(filename)[1].lower()
    if ext == '.pdf':
        pages = _read_pdf(stream)
    elif ext == '.docx':
        pages = _read_docx(stream)
    elif ext == '.txt':
        pages = _read_txt(stream)
    else:
        raise ValueError('Invalid format or corrupted file')

    doc = Document(filename=filename, business_id=business_id)
    db.session.add(doc)
    db.session.commit()

    texts = []
    metadatas = []
    for i, page in enumerate(pages, start=1):
        if not page:
            continue
        for chunk in _chunk_text(page):
            texts.append(chunk)
            metadatas.append({'doc_id': doc.id, 'business_id': business_id, 'page': i, 'text': chunk})

    if texts:
        vector_store = _get_vector_store()
        if vector_store is not None:
            vector_store.add(texts, metadatas)
        else:
            _persist_chunks(texts, metadatas)
    return {'doc_id': doc.id, 'chunks': len(texts)}
